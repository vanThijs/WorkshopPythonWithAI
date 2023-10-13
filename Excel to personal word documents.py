import pandas as pd
from docx import Document

# Read data from the Excel file  
excel_file = "C:\\Users\\lucg7043\\Documents\\Python\\Temp\\Student_Test\\student_data.xlsx"  # Replace with the actual Excel file path
#excel_file = "student_data.xlsx"  # Replace with the actual Excel file path
df = pd.read_excel(excel_file)

# Iterate through the rows and generate Word files
for index, row in df.iterrows():
    first_name = row["first name"]
    last_name = row["last name"]
    grade = row["grade"]
    comment = row["comment"]
    teacher_name = row["teacher name"]

    # Create a Word document
    doc = Document()
    doc.add_paragraph(f"Dear {first_name} {last_name},")
    doc.add_paragraph(f"You have scored {grade}/20.")
    doc.add_paragraph(f"The teaching corps had the following remarks: {comment}.")
    doc.add_paragraph(f"Have a nice holiday, {teacher_name}.")

    # Save the Word document with the first and last name of the student
    output_file = f"{first_name}_{last_name}.docx"
    doc.save(output_file)

    print(f"Word file '{output_file}' generated for {first_name} {last_name}")
    
    
    '''
ChatGPT command: 
write python code which opens an excel file, with data organised in the following columns:
"first name", "last name", "grade", ''comment'', "teacher name"

A word file should be generated from this data, containing the following information:
"Dear {first name} {last name}, You have scored {grade}/20. The teaching corps had the following remarks: {comment}.  Have a nice holliday, {teacher name}."
This file should be saved with the first and lastname of the student in the filename.
'''
